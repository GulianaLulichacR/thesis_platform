import io
from pathlib import Path

import aiofiles
import docx

from app.core.exceptions import DocumentLoadError
from app.core.logging import get_logger

logger = get_logger(__name__)


class DOCXLoader:
    """Extracts text and metadata from DOCX files."""

    async def load_from_path(self, path: Path) -> dict:
        try:
            async with aiofiles.open(path, "rb") as f:
                content = await f.read()
            return self._parse(content, filename=path.name)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to load DOCX '{path.name}': {exc}") from exc

    async def load_from_bytes(self, data: bytes, filename: str = "document.docx") -> dict:
        try:
            return self._parse(data, filename=filename)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to parse DOCX '{filename}': {exc}") from exc

    def _parse(self, data: bytes, filename: str) -> dict:
        document = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Extract headings as section hints
        headings = [
            p.text for p in document.paragraphs if p.style.name.startswith("Heading")
        ]

        core_props = document.core_properties

        return {
            "filename": filename,
            "pages": None,  # DOCX has no direct page count without rendering
            "full_text": full_text,
            "paragraphs": paragraphs,
            "headings": headings,
            "size_bytes": len(data),
            "title": core_props.title,
            "author": core_props.author,
            "word_count": len(full_text.split()),
        }
